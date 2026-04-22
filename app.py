import streamlit as st
import os
import time
import requests
from io import BytesIO
from docx import Document
import tempfile

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_pinecone import PineconeVectorStore
from langchain_groq import ChatGroq
from langchain_core.embeddings import Embeddings

# --- 1. SECRETS & CONFIGURATION ---
os.environ["PINECONE_API_KEY"] = st.secrets["PINECONE_API_KEY"]
os.environ["GROQ_API_KEY"] = st.secrets["GROQ_API_KEY"]
hf_token = st.secrets["HUGGINGFACEHUB_API_TOKEN"]

index_name = "audit-db"

# --- 2. CUSTOM BULLETPROOF EMBEDDER ---
# This forces the free API to wake up and catches blank errors
# --- 2. CUSTOM BULLETPROOF EMBEDDER ---
# This forces the free API to wake up and catches blank errors
class SafeHFEmbeddings(Embeddings):
    def __init__(self, api_key, model):
        # Updated to the correct Hugging Face API routing
        self.api_url = f"https://router.huggingface.co/hf-inference/models/{model}/pipeline/feature-extraction"
        self.headers = {"Authorization": f"Bearer {api_key}"}

    def embed_documents(self, texts):
        for attempt in range(3):
            # We explicitly tell Hugging Face to wake the model up and wait
            payload = {"inputs": texts, "options": {"wait_for_model": True}}
            response = requests.post(self.api_url, headers=self.headers, json=payload)
            
            if response.status_code == 200:
                try:
                    return response.json()
                except Exception:
                    raise Exception(f"API returned invalid data. Raw response: {response.text}")
            elif response.status_code == 503:
                # Model is asleep. Wait 5 seconds and try again.
                time.sleep(5)
                continue
            else:
                raise Exception(f"HF Error {response.status_code}: {response.text}")
                
        raise Exception("Hugging Face API timed out. The server is currently too busy.")

    def embed_query(self, text):
        return self.embed_documents([text])[0]
# --- 3. MAIN APPLICATION ---
st.set_page_config(page_title="Cloud Audit AI", layout="centered")
st.title("🛡️ Cloud Audit Engine")
st.markdown("Public AI auditing tool powered by Groq and Pinecone.")

# Initialize Cloud AI Models
llm = ChatGroq(model_name="llama3-8b-8192")

# Use our new bulletproof embedder
embeddings = SafeHFEmbeddings(
    api_key=hf_token, 
    model="sentence-transformers/all-MiniLM-L6-v2"
)

vector_store = PineconeVectorStore(index_name=index_name, embedding=embeddings)

# --- 4. SIDEBAR: DOCUMENT UPLOAD ---
with st.sidebar:
    st.header("Document Repository")
    uploaded_file = st.file_uploader("Upload Audit PDF", type=["pdf"])
    
    if st.button("Process & Upload to Cloud"):
        if uploaded_file is not None:
            with st.spinner("Processing PDF..."):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(uploaded_file.getvalue())
                    tmp_path = tmp.name
                
                try:
                    loader = PyPDFLoader(tmp_path)
                    docs = loader.load()
                    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
                    chunks = text_splitter.split_documents(docs)
                    
                    for chunk in chunks:
                        chunk.metadata["source"] = uploaded_file.name
                    
                    batch_size = 10
                    progress_text = st.empty()
                    
                    for i in range(0, len(chunks), batch_size):
                        batch = chunks[i : i + batch_size]
                        progress_text.text(f"Uploading batch {i//batch_size + 1}...")
                        vector_store.add_documents(batch)
                        time.sleep(1) 
                        
                    progress_text.empty()
                    st.success(f"Successfully uploaded {uploaded_file.name} to Pinecone!")
                    
                except Exception as e:
                    st.error(f"Upload Failed: {str(e)}")
                finally:
                    os.remove(tmp_path)
        else:
            st.warning("Please upload a file.")

# --- 5. MAIN UI: GENERATE QUESTIONNAIRE ---
focus_area = st.text_input("Audit Focus Area", placeholder="e.g., Data Privacy Protocol")
doc_target = st.text_input("Target Document Name (Exact PDF name uploaded)", placeholder="e.g., policy.pdf")

if st.button("Generate Questionnaire"):
    if focus_area and doc_target:
        with st.spinner("Waking up AI and searching documents..."):
            try:
                retriever = vector_store.as_retriever(
                    search_kwargs={"k": 4, "filter": {"source": doc_target}}
                )
                retrieved_docs = retriever.invoke(focus_area)
                context = "\n\n".join([doc.page_content for doc in retrieved_docs])
                
                if not context:
                    st.warning("No relevant text found. Did you type the exact PDF filename?")
                else:
                    prompt = f"""You are an expert auditor. Based ONLY on the excerpts below from {doc_target}, generate a 5-question audit checklist about: {focus_area}.
                    Excerpts: {context}"""
                    
                    response = llm.invoke(prompt)
                    
                    st.session_state['last_result'] = response.content
                    st.session_state['last_focus'] = focus_area
                    st.markdown("### Generated Questionnaire")
                    st.write(response.content)
            except Exception as e:
                st.error(f"Error generating audit: {str(e)}")

# --- 6. EXPORT TO WORD ---
if 'last_result' in st.session_state:
    st.divider()
    st.subheader("Export")
    
    doc = Document()
    doc.add_heading(f"Audit Questionnaire: {st.session_state.get('last_focus', 'Topic')}", 0)
    doc.add_paragraph(st.session_state['last_result'])
    
    bio = BytesIO()
    doc.save(bio)
    
    st.download_button(
        label="Download as Word (.docx)",
        data=bio.getvalue(),
        file_name="audit_questionnaire.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
